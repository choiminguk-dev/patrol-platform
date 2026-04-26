"""적극행정 보고서 DOCX + PDF 생성"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 여백 설정
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.space_after = Pt(20)

def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.space_before = Pt(15)
    p.space_after = Pt(8)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        # 헤더 배경색
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'E8F5E9',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # 데이터
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()  # 간격

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ===== 본문 작성 =====

add_title('창의행정 추진계획(안)')
add_subtitle('「AI 환경순찰」 순찰 업무 자동화 시스템 구축')

# 추진배경
add_section_header('▣ 추진배경')
add_bullet('환경순찰 후 순찰일지 작성에 1일 평균 2~3시간 소요 (수작업 기록)')
add_bullet('현장사진 정리, 위치 기재, 실적 집계 등 반복적 행정 업무 과다')
add_bullet('공무관·지킴이 등 현장 인력이 사진을 전달하면 담당자가 수동 정리')
add_bullet('순찰 기록의 정확성·적시성 확보 어려움 (사후 기억 의존)')
add_bullet('AI 기술을 활용한 업무 자동화로 행정 효율화 필요성 대두')

# 추진내용
add_section_header('▣ 추진내용')

add_body('1) AI 환경순찰 시스템 구축')
add_table(
    ['구분', '내용'],
    [
        ['시스템명', 'AI 환경순찰 통합 관리 시스템'],
        ['접속주소', 'https://patrol.ai.kr'],
        ['사용대상', '청소담당, 안전담당 (2명)'],
        ['클라우드', 'NHN Cloud\n · 과기정통부·KISA 「클라우드 보안 인증(CSAP)」 취득\n · 국내 리전(Region) 기반 데이터 주권 확보'],
        ['보안', 'HTTPS 암호화, PIN 인증, HttpOnly 세션'],
        ['라이선스', 'MIT License (전 세계 공공·민간 자유 사용 국제 표준 라이선스)'],
    ]
)

add_body('2) 핵심 기능')
add_table(
    ['기능', '설명', '효과'],
    [
        ['사진 일괄 업로드', '현장사진 30장 동시 압축·업로드', '업로드 시간 대폭 단축'],
        ['AI 주소판 OCR', '도로명 표지판 자동 인식', '위치 수기 입력 불필요'],
        ['사진 자동 그룹핑', '주소판+현장사진 패턴 자동 분류', '사진 정리 시간 제거'],
        ['순찰일지 자동 생성', '날짜 선택 → 공문서 형식 즉시 생성', '작성 시간 약 2시간 → 수 분'],
        ['평가 진척도', '7개 평가항목 100점 기준 현황 자동 집계', '실적 관리 자동화'],
    ]
)

add_body('3) 업무 흐름 개선')
add_table(
    ['구분', '기존', '개선'],
    [
        ['사진 전달', '카톡 전송 → PC 다운로드', '시스템 직접 업로드 (1분)'],
        ['위치 기재', '수기 입력 (기억 의존)', 'AI 주소판 자동 인식'],
        ['사진 정리', '폴더별 수동 분류', 'AI 자동 그룹핑'],
        ['순찰일지', '엑셀/한글 수작업 (약 2~3시간)', '자동 생성 (수 분)'],
        ['실적 관리', '수기 집계', '자동 집계 + 진척도 표시'],
    ]
)

# 소요예산
add_section_header('▣ 소요예산')
add_table(
    ['항목', '비용', '비고'],
    [
        ['시스템 개발', '0원', '자체 개발'],
        ['클라우드 서버 (NHN)', '약 52,000원/월', 'CSAP 인증 국내 서버'],
        ['도메인 (patrol.ai.kr)', '16,500원/년', ''],
        ['연간 총 비용', '약 640,000원', '월 약 53,000원'],
    ]
)
add_body('※ 기존 수작업 대비 담당자 업무시간 연간 상당 시간 절감 효과')

# 기대효과
add_section_header('▣ 기대효과')
add_table(
    ['효과', '내용'],
    [
        ['업무시간 절감', '순찰일지 작성 약 2~3시간 → 수 분 (약 90% 이상 절감)'],
        ['기록 정확성', 'AI 주소 인식으로 위치정보 자동화, 누락 방지'],
        ['실적 관리', '평가항목별 진척도 자동 집계, 실시간 현황 파악'],
        ['보안 확보', 'CSAP 인증 국내 클라우드 + HTTPS 암호화'],
        ['확산 가능', '오픈소스(MIT) 공개, 타 동·구 즉시 적용 가능'],
    ]
)

# 추진일정
add_section_header('▣ 추진일정')
add_table(
    ['단계', '내용', '일정'],
    [
        ['1단계', '시스템 설계 및 MVP 개발', '완료 (4.4)'],
        ['2단계', 'AI 주소판 OCR·사진 그룹핑 개발', '완료 (4.4)'],
        ['3단계', 'NHN Cloud 배포 + HTTPS 적용', '완료 (4.5)'],
        ['4단계', '실무 적용 및 데이터 축적', '4월~'],
        ['5단계', '활용도 분석 및 고도화', '6월~'],
    ]
)

# 향후계획
add_section_header('▣ 향후계획')
add_bullet('4~6월 실무 적용을 통한 데이터 축적 및 시스템 안정화')
add_bullet('상반기 청소평가 활용 (평가 진척도 자동 관리)')
add_bullet('사용자 피드백 반영 고도화 (보고서 양식 추가, 지도 연동 등)')
add_bullet('타 동·구 확산 목적의 오픈소스(MIT) 공개')
add_body('  GitHub: https://github.com/choiminguk-dev/patrol-platform')

# 끝
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.space_before = Pt(30)
run = p.add_run('- 끝 -')
run.font.size = Pt(10)
run.font.name = '맑은 고딕'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 저장
output_dir = r'C:\Users\user\Downloads'
docx_path = os.path.join(output_dir, 'AI환경순찰_창의행정_보고서.docx')
doc.save(docx_path)
print(f'DOCX 생성 완료: {docx_path}')

# PDF 변환 시도
try:
    from docx2pdf import convert
    pdf_path = os.path.join(output_dir, 'AI환경순찰_창의행정_보고서.pdf')
    convert(docx_path, pdf_path)
    print(f'PDF 생성 완료: {pdf_path}')
except ImportError:
    print('PDF 변환: docx2pdf 미설치. DOCX를 열어서 "다른 이름으로 저장 → PDF"로 변환하세요.')
except Exception as e:
    print(f'PDF 변환 실패: {e}. DOCX를 열어서 수동 PDF 변환하세요.')
